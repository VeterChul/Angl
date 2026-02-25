from pdf2docx import Converter
from docx import Document
import fitz


def convert_pdf_to_docx(pdf_path, docx_path=None):
    """
    Конвертирует PDF файл в DOCX
    
    Args:
        pdf_path: путь к PDF файлу
        docx_path: путь для сохранения DOCX (если None, то заменяет расширение на .docx)
    
    Returns:
        путь к созданному DOCX файлу
    """
    # Если docx_path не указан, создаем путь с тем же именем, но с расширением .docx
    if docx_path is None:
        docx_path = pdf_path.replace('.pdf', '.docx')
    
    print(f"Конвертируем {pdf_path} в {docx_path}...")
    
    # Создаем конвертер и конвертируем
    cv = Converter(pdf_path)
    cv.convert(docx_path, start=0, end=None)  # start и end - страницы (None = все)
    cv.close()
    
    return docx_path

def find_underlined_text(docx_path):
    """
    Находит весь подчеркнутый текст в DOCX файле
    """
    doc = Document(docx_path)
    
    underlined_texts = []
    
    
    # 1. Проверяем параграфы
    for para_num, paragraph in enumerate(doc.paragraphs, 1):
        for run in paragraph.runs:
            if run.underline and run.text.strip():
                underlined_texts.append({
                    "type": "paragraph",
                    "location": f"Параграф {para_num}",
                    "text": run.text,
                    "bold": run.bold,
                    "italic": run.italic,
                    "full_paragraph": paragraph.text[:100] + "..." if len(paragraph.text) > 100 else paragraph.text
                })
    
    # 2. Проверяем таблицы
    for table_num, table in enumerate(doc.tables, 1):
        for row_num, row in enumerate(table.rows, 1):
            for cell_num, cell in enumerate(row.cells, 1):
                for para_num, paragraph in enumerate(cell.paragraphs, 1):
                    for run in paragraph.runs:
                        if run.underline and run.text.strip():
                            underlined_texts.append({
                                "type": "table",
                                "location": f"Таблица {table_num}, Ячейка [{row_num},{cell_num}]",
                                "text": run.text,
                                "bold": run.bold,
                                "italic": run.italic,
                                "full_paragraph": paragraph.text[:100] + "..." if len(paragraph.text) > 100 else paragraph.text
                            })
    
    # Выводим результаты
    
    ret = []

    for item in underlined_texts:
        ret.append(item['text'])
        
        # print(f"📍 {item['location']}:")
        # print(f"   Текст: '{item['text']}'")
        # print(f"   Стиль: жирный={item['bold']}, курсив={item['italic']}")
        # print()
    
    return set(ret)

# Использование
#docx_file = convert_pdf_to_docx("fold/03/pdf/ans.pdf")
#print(find_underlined_text(docx_file))
#diagnose_pdf("fold/03/pdf/ans.pdf")
